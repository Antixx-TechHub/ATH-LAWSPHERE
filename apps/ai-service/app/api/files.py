"""
File processing API endpoints.
"""

from typing import Optional, List
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, BackgroundTasks, Depends
from pydantic import BaseModel
from enum import Enum
import uuid
import io
from pathlib import Path
import json

# Import for text extraction
try:
    from pypdf import PdfReader
    HAS_PDF_SUPPORT = True
except ImportError:
    try:
        import PyPDF2
        from PyPDF2 import PdfReader
        HAS_PDF_SUPPORT = True
    except ImportError:
        HAS_PDF_SUPPORT = False

try:
    from docx import Document
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

# OCR support for images
try:
    from PIL import Image
    import pytesseract
    HAS_OCR_SUPPORT = True
except ImportError:
    HAS_OCR_SUPPORT = False

from app.routing.privacy_scanner import PrivacyScanner
from app.routing.audit_logger import AuditLogger
from app.db import get_db
from app.services import session_store
from sqlalchemy.ext.asyncio import AsyncSession

router = APIRouter()
privacy_scanner = PrivacyScanner()
audit_logger = AuditLogger()


class FileStatus(str, Enum):
    UPLOADING = "uploading"
    PROCESSING = "processing"
    READY = "ready"
    ERROR = "error"


class FileMetadata(BaseModel):
    id: str
    filename: str
    original_name: str
    mime_type: str
    size: int
    status: FileStatus
    extracted_text: Optional[str] = None
    metadata: Optional[dict] = None
    is_sensitive: Optional[bool] = None
    pii_detected: Optional[bool] = None
    sensitivity_score: Optional[float] = None


class ProcessingResult(BaseModel):
    file_id: str
    status: FileStatus
    extracted_text: Optional[str] = None
    pages: Optional[int] = None
    entities: Optional[List[dict]] = None
    summary: Optional[str] = None


@router.post("/upload", response_model=FileMetadata)
async def upload_file(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
    user_id: Optional[str] = Form(None),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: AsyncSession = Depends(get_db),
):
    """
    Upload a file for processing.
    Supports PDF, DOC, DOCX, images, audio, and video files.
    """
    # Validate file type
    allowed_types = [
        "application/pdf",
        "application/msword",  # .doc
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
        "application/vnd.ms-excel",  # .xls
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # .xlsx
        "image/png",
        "image/jpeg",
        "image/jpg",
        "image/webp",
        "image/gif",
        "image/tiff",
        "image/bmp",
        "audio/mpeg",
        "audio/wav",
        "video/mp4",
        "text/plain",
        "text/markdown",
        "text/csv",
        "application/rtf",  # RTF files
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file.content_type}",
        )
    
    # Generate file ID
    file_id = str(uuid.uuid4())
    
    # Read file content
    content = await file.read()
    
    # Persist file metadata AND raw content tied to the session
    session_id = await session_store.ensure_session(db, session_id, user_id)
    await session_store.add_file(
        db,
        session_id=session_id,
        name=file.filename or "uploaded_file",
        mime_type=file.content_type,
        size=len(content),
        status=FileStatus.PROCESSING.value,
        file_id=file_id,
        raw_content=content,  # Store the raw file content
    )
    await db.commit()
    
    # Add background processing task (pass db to update extracted text)
    background_tasks.add_task(process_file, file_id, content, file.content_type)
    
    return FileMetadata(
        id=file_id,
        filename=f"{file_id}_{file.filename}",
        original_name=file.filename or "unknown",
        mime_type=file.content_type or "application/octet-stream",
        size=len(content),
        status=FileStatus.PROCESSING,
    )


async def process_file(file_id: str, content: bytes, mime_type: str):
    """Background task to process uploaded file and save extracted text."""
    from app.db import SessionLocal
    from sqlalchemy import update
    import traceback

    try:
        extracted_text = ""
        print(f"[FILE_PROCESS] Starting processing for file {file_id} with mime_type: {mime_type}")

        # Extract text based on file type
        if mime_type == "application/pdf" and HAS_PDF_SUPPORT:
            print(f"[FILE_PROCESS] Extracting PDF text for {file_id}")
            extracted_text = extract_text_from_pdf(content)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" and HAS_DOCX_SUPPORT:
            print(f"[FILE_PROCESS] Extracting DOCX text for {file_id}")
            extracted_text = extract_text_from_docx(content)
        elif mime_type == "application/msword":
            # Legacy .doc files - try to extract what we can
            print(f"[FILE_PROCESS] Attempting legacy DOC extraction for {file_id}")
            extracted_text = extract_text_from_doc(content)
        elif mime_type in ["text/plain", "text/markdown", "text/csv"]:
            print(f"[FILE_PROCESS] Decoding text file {file_id}")
            extracted_text = content.decode('utf-8', errors='ignore')
            print(f"[FILE_PROCESS] Decoded {len(extracted_text)} chars from text file")
        elif mime_type.startswith("image/") and HAS_OCR_SUPPORT:
            print(f"[FILE_PROCESS] Running OCR on image {file_id}")
            extracted_text = extract_text_from_image(content)
        else:
            # For unsupported types, just mark as ready
            extracted_text = "[Binary file - no text extraction available]"

        print(f"[FILE_PROCESS] Extracted {len(extracted_text)} chars from {file_id}")
        if extracted_text:
            print(f"[FILE_PROCESS] Content preview: {extracted_text[:200]}...")

        is_sensitive = False
        pii_detected = False

        # Scan for sensitive content
        if extracted_text and not extracted_text.startswith("["):
            scan_result = privacy_scanner.scan(extracted_text)
            # ScanResult is a dataclass - use attribute access
            is_sensitive = scan_result.force_local
            pii_detected = len(scan_result.pii_found) > 0

            print(f"[FILE_PROCESS] File {file_id} scan results:")
            print(f"[FILE_PROCESS]   - Sensitive: {is_sensitive}")
            print(f"[FILE_PROCESS]   - PII Detected: {pii_detected}")
            print(f"[FILE_PROCESS]   - Sensitivity Score: {scan_result.confidence_score}")

        # Update database with extracted text and processing results
        print(f"[FILE_PROCESS] Updating database for file {file_id}...")
        
        if SessionLocal is None:
            print(f"[FILE_PROCESS] ERROR: SessionLocal is None - database not initialized!")
            return
            
        async with SessionLocal() as db:
            from app.db_models import SessionFile
            stmt = (
                update(SessionFile)
                .where(SessionFile.id == file_id)
                .values(
                    extracted_text=extracted_text,
                    status="ready",
                    is_sensitive=is_sensitive,
                    pii_detected=pii_detected,
                )
            )
            result = await db.execute(stmt)
            await db.commit()
            print(f"[FILE_PROCESS] File {file_id} saved to database - rows affected: {result.rowcount}")

    except Exception as e:
        # Log error and update status to error
        print(f"[FILE_PROCESS] ERROR processing file {file_id}: {str(e)}")
        print(f"[FILE_PROCESS] Traceback: {traceback.format_exc()}")
        try:
            if SessionLocal is not None:
                async with SessionLocal() as db:
                    from app.db_models import SessionFile
                    stmt = (
                        update(SessionFile)
                        .where(SessionFile.id == file_id)
                        .values(status="error")
                    )
                    await db.execute(stmt)
                    await db.commit()
                    print(f"[FILE_PROCESS] Updated file {file_id} status to error")
        except Exception as db_error:
            print(f"[FILE_PROCESS] Failed to update error status: {str(db_error)}")

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF file."""
    try:
        pdf_file = io.BytesIO(content)
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip() if text else "[PDF content could not be extracted - may be scanned/image-based]"
    except Exception as e:
        print(f"Error extracting PDF text: {str(e)}")
        return f"[Error extracting PDF: {str(e)}]"


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts).strip() if text_parts else "[DOCX file appears to be empty]"
    except Exception as e:
        print(f"Error extracting DOCX text: {str(e)}")
        return f"[Error extracting DOCX: {str(e)}]"


def extract_text_from_doc(content: bytes) -> str:
    """Extract text from legacy DOC file (best effort)."""
    try:
        # Try to extract plain text from binary DOC
        # This is a simple approach - binary .doc files have text interspersed
        text = ""
        try:
            # Try decoding as text (works for some DOC files)
            decoded = content.decode('utf-8', errors='ignore')
            # Filter out non-printable characters
            text = ''.join(c for c in decoded if c.isprintable() or c in '\n\r\t ')
            # Clean up excessive whitespace
            import re
            text = re.sub(r'\s+', ' ', text).strip()
        except:
            pass
        
        if len(text) > 100:
            return text
        return "[Legacy DOC format - upload as DOCX for full text extraction]"
    except Exception as e:
        print(f"Error extracting DOC text: {str(e)}")
        return f"[Error extracting DOC: {str(e)}]"


def extract_text_from_image(content: bytes) -> str:
    """Extract text from image using OCR (Tesseract)."""
    try:
        image = Image.open(io.BytesIO(content))
        
        # Convert to RGB if necessary (for PNG with transparency, etc.)
        if image.mode in ('RGBA', 'P'):
            image = image.convert('RGB')
        
        # Run OCR
        text = pytesseract.image_to_string(image, lang='eng')
        
        if text.strip():
            return text.strip()
        return "[No text detected in image - may be a photo or graphic without readable text]"
    except Exception as e:
        print(f"Error extracting text from image: {str(e)}")
        return f"[Error running OCR: {str(e)}]"


@router.get("/{file_id}", response_model=FileMetadata)
async def get_file(file_id: str, db: AsyncSession = Depends(get_db)):
    """Get file metadata and processing status."""
    file_data = await session_store.get_file(db, file_id)
    if not file_data:
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileMetadata(
        id=file_data["id"],
        filename=file_data["name"],
        original_name=file_data["name"],
        mime_type=file_data.get("mime_type", "application/octet-stream"),
        size=file_data.get("size", 0),
        status=FileStatus(file_data.get("status", "ready")),
        extracted_text=file_data.get("extracted_text"),
        metadata=file_data.get("metadata"),
        is_sensitive=file_data.get("is_sensitive"),
        pii_detected=file_data.get("pii_detected"),
        sensitivity_score=file_data.get("sensitivity_score"),
    )


@router.get("/{file_id}/content")
async def get_file_content(file_id: str, db: AsyncSession = Depends(get_db)):
    """Get extracted text content from a processed file."""
    file_data = await session_store.get_file(db, file_id)
    if not file_data:
        raise HTTPException(status_code=404, detail="File not found")
    
    return {
        "file_id": file_id,
        "filename": file_data["name"],
        "content": file_data.get("extracted_text", ""),
        "mime_type": file_data.get("mime_type"),
    }


@router.get("/{file_id}/raw")
async def get_file_raw(file_id: str, db: AsyncSession = Depends(get_db)):
    """Get raw file content (binary) for viewing images, PDFs, etc."""
    from fastapi.responses import Response
    
    file_data = await session_store.get_file(db, file_id)
    if not file_data:
        raise HTTPException(status_code=404, detail="File not found")
    
    raw_content = file_data.get("raw_content")
    if not raw_content:
        raise HTTPException(status_code=404, detail="Raw content not available")
    
    mime_type = file_data.get("mime_type", "application/octet-stream")
    
    return Response(
        content=raw_content,
        media_type=mime_type,
        headers={
            "Content-Disposition": f'inline; filename="{file_data["name"]}"',
            "Cache-Control": "public, max-age=3600",
        }
    )


@router.post("/{file_id}/ocr", response_model=ProcessingResult)
async def trigger_ocr(file_id: str):
    """Trigger OCR processing for an image or scanned document."""
    # TODO: Implement OCR processing
    raise HTTPException(status_code=501, detail="Not implemented")


@router.post("/{file_id}/embed")
async def generate_embeddings(file_id: str):
    """Generate vector embeddings for a processed file."""
    # TODO: Implement embedding generation
    raise HTTPException(status_code=501, detail="Not implemented")


@router.delete("/{file_id}")
async def delete_file(file_id: str, db: AsyncSession = Depends(get_db)):
    """Delete a file and its associated data."""
    deleted = await session_store.delete_file(db, file_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="File not found")
    await db.commit()
    return {"status": "deleted", "file_id": file_id}
